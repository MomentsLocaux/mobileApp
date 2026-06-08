#!/usr/bin/env python3
"""Generate Moments Locaux social login & mobile config Word document."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path

OUTPUT = Path(__file__).parent / "Moments-Locaux-Login-Social-Configuration.docx"


def add_title(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_p(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def build() -> None:
    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title(doc, "Moments Locaux — Guide de configuration")
    sub = doc.add_paragraph("Login social, identifiants mobile, EAS & stores")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Document généré pour le projet mobileApp (Expo / React Native)")
    doc.add_paragraph("Date de référence : juin 2026")
    doc.add_page_break()

    # 1. Identifiants
    add_h(doc, "1. Identifiants de l'application")
    add_p(
        doc,
        "Ces valeurs sont l'ADN technique de l'app. Le bundle ID et le scheme "
        "ne doivent pas changer après la beta (sinon = nouvelle app côté stores).",
    )
    add_table(
        doc,
        ["Élément", "Valeur", "Où c'est défini", "Rôle"],
        [
            ["Nom affiché", "Moments Locaux", "app.config.ts → name", "Nom sous l'icône / dans l'app"],
            ["Slug Expo", "moments-locaux", "app.config.ts → slug", "Identifiant projet Expo"],
            ["Bundle ID iOS", "com.momentslocs.app", "app.config.ts → ios.bundleIdentifier", "Identifiant unique App Store / Apple"],
            ["Package Android", "com.momentslocs.app", "app.config.ts → android.package", "Identifiant unique Play Store"],
            ["URL scheme (deep link)", "momentslocaux", "app.config.ts → scheme", "Retour OAuth dans l'app"],
            ["Redirect app OAuth", "momentslocaux://auth/callback", "oauth.service.ts", "URL de retour après login social"],
            ["Version", "1.0.0", "app.config.ts", "Version marketing"],
            ["iOS buildNumber", "1", "app.config.ts", "Numéro de build iOS"],
            ["Android versionCode", "1", "app.config.ts", "Numéro de build Android"],
        ],
    )

    add_h(doc, "1.1 Changer le nom plus tard ?", level=2)
    add_bullets(
        doc,
        [
            "Nom affiché (« Moments Locaux ») : simple, nécessite un nouveau build.",
            "Titre App Store / Play Store : modifiable dans les consoles store.",
            "Bundle ID / package / scheme : à NE PAS changer après publication (nouvelle app, OAuth à reconfigurer).",
        ],
    )

    doc.add_page_break()

    # 2. Architecture
    add_h(doc, "2. Architecture du login social")
    add_p(doc, "L'app utilise Supabase Auth comme backend d'authentification.")
    add_h(doc, "2.1 Flux général (Google & Facebook)", level=2)
    add_bullets(
        doc,
        [
            "L'utilisateur appuie sur Google ou Facebook dans l'app.",
            "Un navigateur in-app s'ouvre (expo-web-browser, PKCE).",
            "Le provider redirige vers Supabase : https://prymkgkafaovhzopslea.supabase.co/auth/v1/callback",
            "Supabase crée la session et redirige vers l'app : momentslocaux://auth/callback",
            "L'app échange le code OAuth contre une session Supabase.",
            "Redirection vers la carte ou l'onboarding selon le profil.",
        ],
    )
    add_h(doc, "2.2 Flux Apple (iOS)", level=2)
    add_bullets(
        doc,
        [
            "Sur iPhone : Sign In with Apple NATIF (expo-apple-authentication).",
            "Jeton identityToken envoyé à Supabase via signInWithIdToken.",
            "Nécessite la capability Apple Sign In + compte Apple Developer payant.",
            "Sur Android / fallback : flux navigateur OAuth comme Google.",
        ],
    )
    add_h(doc, "2.3 Fichiers code concernés", level=2)
    add_table(
        doc,
        ["Fichier", "Rôle"],
        [
            ["src/services/oauth.service.ts", "Logique OAuth (browser + Apple natif)"],
            ["src/services/auth.service.ts", "signInWithProvider, finalisation session"],
            ["src/hooks/useAuth.ts", "Hook exposé aux écrans"],
            ["src/components/auth/SocialLoginButtons.tsx", "Boutons Google / Apple / Facebook"],
            ["src/screens/auth/LoginScreen.tsx", "Écran connexion"],
            ["src/screens/auth/RegisterScreen.tsx", "Écran inscription + consentement CGU"],
            ["app/auth/callback.tsx", "Gestion deep link au cold start"],
        ],
    )

    doc.add_page_break()

    # 3. Supabase
    add_h(doc, "3. Configuration Supabase")
    add_table(
        doc,
        ["Paramètre", "Valeur"],
        [
            ["URL projet", "https://prymkgkafaovhzopslea.supabase.co"],
            ["Callback OAuth (providers)", "https://prymkgkafaovhzopslea.supabase.co/auth/v1/callback"],
            ["Site URL (dashboard Auth)", "momentslocaux://auth/callback (ou URL de prod)"],
        ],
    )
    add_h(doc, "3.1 Redirect URLs (Authentication → URL Configuration)", level=2)
    add_bullets(
        doc,
        [
            "momentslocaux://auth/callback",
            "momentslocaux://**",
        ],
    )
    add_p(doc, "Important : avec un dev build Expo, pas besoin d'URLs exp:// (Expo Go).")

    doc.add_page_break()

    # 4. Google
    add_h(doc, "4. Google OAuth")
    add_h(doc, "4.1 Google Cloud Console", level=2)
    add_bullets(
        doc,
        [
            "Créer un projet Google Cloud (ou utiliser l'existant).",
            "APIs & Services → Credentials → Create OAuth client ID.",
            "Type : Web application (pas iOS/Android pour ce flux Supabase).",
            "Authorized redirect URI : https://prymkgkafaovhzopslea.supabase.co/auth/v1/callback",
        ],
    )
    add_h(doc, "4.2 Supabase → Auth → Providers → Google", level=2)
    add_bullets(
        doc,
        [
            "Activer le provider Google.",
            "Renseigner Client ID + Client Secret (client Web Google).",
        ],
    )

    # 5. Apple
    add_h(doc, "5. Apple Sign In")
    add_h(doc, "5.1 Prérequis", level=2)
    add_bullets(
        doc,
        [
            "Apple Developer Program payant (99 €/an) — Personal Team insuffisant.",
            "App ID com.momentslocs.app avec capability Sign In with Apple activée.",
            "Entitlement natif : com.apple.developer.applesignin (ajouté par expo prebuild).",
        ],
    )
    add_h(doc, "5.2 Apple Developer Portal", level=2)
    add_bullets(
        doc,
        [
            "Identifiers → com.momentslocs.app → Sign In with Apple → Enable.",
            "Keys → créer une clé .p8 pour Sign In with Apple (si secret Supabase requis).",
            "Team ID, Key ID, Bundle ID à noter pour Supabase.",
        ],
    )
    add_h(doc, "5.3 Supabase → Auth → Providers → Apple", level=2)
    add_bullets(
        doc,
        [
            "Activer Apple provider.",
            "Client IDs : com.momentslocs.app",
            "Secret : généré via la clé .p8 Apple (Services ID si flux web, ici natif iOS).",
        ],
    )

    doc.add_page_break()

    # 6. Facebook
    add_h(doc, "6. Facebook / Meta")
    add_table(
        doc,
        ["Paramètre", "Valeur"],
        [
            ["Meta App ID", "1310925674348998"],
            ["Bundle ID iOS", "com.momentslocs.app"],
            ["Package Android", "com.momentslocs.app"],
            ["Key hash Android (debug)", "Xo8WBi6jzSxKDVR4drqm84yr9iU="],
            ["SHA-1 debug keystore", "5E:8F:16:06:2E:A3:CD:2C:4A:0D:54:78:76:BA:A6:F3:8C:AB:F6:25"],
        ],
    )
    add_h(doc, "6.1 Où configurer dans Meta (PAS sous Facebook Login → Settings)", level=2)
    add_bullets(
        doc,
        [
            "Paramètres de l'app → De base → tout en bas → + Ajouter une plateforme.",
            "iOS : Bundle ID com.momentslocs.app (Store ID vide si pas encore sur l'App Store).",
            "Android : package com.momentslocs.app + key hash debug (et release plus tard).",
            "Plateforme Site web : NON nécessaire pour l'app mobile.",
        ],
    )
    add_h(doc, "6.2 Facebook Login → Paramètres (OAuth web)", level=2)
    add_bullets(
        doc,
        [
            "URI de redirection OAuth valides : https://prymkgkafaovhzopslea.supabase.co/auth/v1/callback",
            "Attention orthographe projet : prymkgkafaovhzopslea (avec le k).",
            "Mode Développement : ajouter les comptes Facebook comme testeurs de l'app.",
        ],
    )
    add_h(doc, "6.3 Supabase → Auth → Providers → Facebook", level=2)
    add_bullets(doc, ["App ID + App Secret depuis Meta → Paramètres → De base."])

    add_h(doc, "6.4 Pages légales (exigées par Meta)", level=2)
    add_table(
        doc,
        ["Page", "URL GitHub Pages"],
        [
            ["Politique de confidentialité", "https://momentslocaux.github.io/mobileApp/privacy.html"],
            ["Suppression des données", "https://momentslocaux.github.io/mobileApp/delete.html"],
            ["CGU", "https://momentslocaux.github.io/mobileApp/terms.html"],
        ],
    )

    doc.add_page_break()

    # 7. EAS
    add_h(doc, "7. Configuration EAS Build")
    add_p(doc, "Fichier eas.json — profils disponibles :")
    add_table(
        doc,
        ["Profil", "Usage", "Particularités"],
        [
            ["development", "Dev build sur iPhone/Android test", "developmentClient: true, distribution internal"],
            ["preview", "Build interne staging", "iOS simulateur, Android APK"],
            ["production", "Soumission stores", "autoIncrement build numbers"],
        ],
    )
    add_h(doc, "7.1 Commandes essentielles", level=2)
    add_bullets(
        doc,
        [
            "npx eas-cli login",
            "npx eas-cli init  →  récupérer EAS_PROJECT_ID dans .env",
            "npx eas-cli credentials --platform ios",
            "npx eas-cli build --profile development --platform ios --clear-cache",
            "npx expo start --dev-client  (après installation du build)",
        ],
    )
    add_h(doc, "7.2 Variables d'environnement (.env)", level=2)
    add_table(
        doc,
        ["Variable", "Rôle"],
        [
            ["EXPO_PUBLIC_SUPABASE_URL", "URL API Supabase"],
            ["EXPO_PUBLIC_SUPABASE_ANON_KEY", "Clé publique anon Supabase"],
            ["EXPO_PUBLIC_MAPBOX_TOKEN", "Token carte Mapbox"],
            ["EAS_PROJECT_ID", "ID projet Expo (push tokens + EAS builds)"],
        ],
    )
    add_p(doc, "Ne jamais committer de secrets (service role, .p8, App Secret) dans git.")

    doc.add_page_break()

    # 8. Apple capabilities
    add_h(doc, "8. Apple Developer — Capabilities & erreurs Xcode")
    add_p(
        doc,
        "Erreur typique avec Personal Team (compte gratuit) :",
        bold=True,
    )
    add_p(
        doc,
        "« Personal development teams do not support Sign In with Apple and Push Notifications »",
    )
    add_h(doc, "8.1 Capabilities requises sur com.momentslocs.app", level=2)
    add_table(
        doc,
        ["Capability", "Pourquoi", "Entitlement"],
        [
            ["Sign In with Apple", "Login Apple natif iOS", "com.apple.developer.applesignin"],
            ["Push Notifications", "Notifications push Expo", "aps-environment = development/production"],
        ],
    )
    add_h(doc, "8.2 Solution recommandée (Option A)", level=2)
    add_bullets(
        doc,
        [
            "S'inscrire au Apple Developer Program payant.",
            "Activer Push + Sign In with Apple sur l'App ID.",
            "Builder via EAS (EAS gère certificats et provisioning profiles).",
            "Uploader clé APNs .p8 dans Expo pour les push.",
            "Supprimer l'ancienne app du téléphone avant d'installer le nouveau build.",
        ],
    )

    add_h(doc, "8.3 Icône de l'app", level=2)
    add_bullets(
        doc,
        [
            "L'icône écran d'accueil est figée dans le binaire natif (pas via expo start).",
            "Source : assets/images/icon.png (1024×1024).",
            "Après changement d'icône : npx expo prebuild --no-install puis eas build.",
            "Fichiers natifs : ios/.../AppIcon.appiconset/ et android/.../ic_launcher*.webp",
        ],
    )

    doc.add_page_break()

    # 9. Android
    add_h(doc, "9. Android — points spécifiques")
    add_bullets(
        doc,
        [
            "Key hash debug : Xo8WBi6jzSxKDVR4drqm84yr9iU= (keystore debug Expo/React Native).",
            "Key hash release : à générer depuis le keystore de production EAS au moment du release.",
            "Commande key hash : keytool -exportcert -alias androiddebugkey -keystore ~/.android/debug.keystore -storepass android -keypass android | openssl sha1 -binary | openssl base64",
            "OpenJDK sur Mac (si besoin) : /opt/homebrew/opt/openjdk@17/bin/keytool",
        ],
    )

    # 10. Test checklist
    add_h(doc, "10. Checklist de test (iPhone)")
    add_table(
        doc,
        ["Étape", "Vérification", "Succès attendu"],
        [
            ["1. Apple", "Bouton Apple → Face ID", "Session créée, provider apple dans Supabase Users"],
            ["2. Google", "Bouton Google → compte", "Retour app → carte ou onboarding"],
            ["3. Facebook", "Compte testeur Meta en mode Dev", "Retour app → carte ou onboarding"],
            ["4. Deep link", "Redirect URLs Supabase OK", "Pas de retour bloqué au login"],
            ["5. Onboarding", "Premier login sans profil complet", "Écran 4 étapes puis carte"],
            ["6. Push", "Après rebuild avec expo-notifications", "Permission + token dans device_push_tokens"],
        ],
    )

    add_h(doc, "11. Erreurs fréquentes", level=1)
    add_table(
        doc,
        ["Erreur", "Cause", "Solution"],
        [
            ["ExpoPushTokenManager not found", "Dev build sans module natif push", "eas build development ou fix lazy-load"],
            ["redirect_uri_mismatch", "URI Supabase absente chez Google/Meta", "Ajouter callback Supabase"],
            ["Facebook App not available", "Compte pas testeur en mode Dev", "Ajouter testeur Meta"],
            ["Apple erreur immédiate", "Personal Team ou provider mal configuré", "Compte payant + Supabase Apple"],
            ["Vieille icône", "ios/AppIcon pas resynchronisé", "expo prebuild + eas build --clear-cache"],
            ["Xcode provisioning errors", "Personal Team + entitlements", "Compte payant ou EAS build"],
        ],
    )

    add_h(doc, "12. Références utiles", level=1)
    add_bullets(
        doc,
        [
            "Supabase Auth docs : https://supabase.com/docs/guides/auth",
            "Expo EAS Build : https://docs.expo.dev/build/introduction/",
            "Meta plateformes : developers.facebook.com → Paramètres app → De base",
            "Apple Developer : https://developer.apple.com/account",
            "Repo mobileApp : https://github.com/MomentsLocaux/mobileApp",
        ],
    )

    doc.add_paragraph()
    footer = doc.add_paragraph("— Fin du document —")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    build()
